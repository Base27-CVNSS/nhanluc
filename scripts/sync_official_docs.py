#!/usr/bin/env python3
"""Synchronize the 18 public dossier files from Vinh Long's paperless-meeting portal.

The script prefers direct links exposed in the rendered DOM. If the attachment list is
collapsed, Playwright clicks the agenda entry containing 617/TTr-UBND first. Files are
saved with stable 01-18 names and plain-text copies are produced for GitHub search.
"""
from __future__ import annotations

import hashlib
import json
import re
import sys
import unicodedata
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from urllib.parse import urljoin

from docx import Document
from playwright.sync_api import Page, TimeoutError as PlaywrightTimeoutError, sync_playwright
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
FILES_DIR = ROOT / "files"
TEXT_DIR = ROOT / "text"
REPORT_PATH = ROOT / "sync_report.json"
MEETING_URL = "https://hopkhonggiay.vinhlong.dcs.vn/Pages/Qr.aspx?Id=16834"


@dataclass(frozen=True)
class Spec:
    index: int
    source_name: str
    output_name: str


SPECS = [
    Spec(1, "1. To trinh UBND tinh trinh du thao Nghi quyet.pdf", "01_To_trinh_UBND_tinh_trinh_du_thao_Nghi_quyet.pdf"),
    Spec(2, "2. Du thao Nghi quyet -trinh thong qua HDND tinh chuan (1).docx", "02_Du_thao_Nghi_quyet_trinh_thong_qua_HDND_tinh.docx"),
    Spec(3, "3. BC-tong ket cong tac dao tao va cac chinh sach giai doan 2021-2025.pdf", "03_BC_tong_ket_cong_tac_dao_tao_chinh_sach_2021_2025.pdf"),
    Spec(4, "4. Bang so sanh thuyet minh chinh sach.pdf", "04_Bang_so_sanh_thuyet_minh_chinh_sach.pdf"),
    Spec(5, "4. BC-tong hop y kienn gop y lan 1.signed.pdf", "05_BC_tong_hop_y_kien_gop_y_lan_1.pdf"),
    Spec(6, "4. BC-tong hop y kienn gop y lan2.signed.pdf", "06_BC_tong_hop_y_kien_gop_y_lan_2.pdf"),
    Spec(7, "4.2. BC-tong hop y kienn gop y lan2.signed.pdf", "07_4_2_BC_tong_hop_y_kien_gop_y_lan_2.pdf"),
    Spec(8, "05- BB.pdf", "08_05_BB_Hop_thong_nhat_noi_dung.pdf"),
    Spec(9, "5. Bc danh gia tac dong chinh sach dao tao nhan luc (1).pdf", "09_BC_danh_gia_tac_dong_chinh_sach_dao_tao_nhan_luc.pdf"),
    Spec(10, "5.3.BC-tong hop gop y lan 3 - chinh sua.signed (1).pdf", "10_BC_tong_hop_gop_y_lan_3_chinh_sua.pdf"),
    Spec(11, "7. Bao cao Tham dinh NQ quy dinh chinh sach dao tao- boi duong.signed.signed.signed.pdf", "11_Bao_cao_Tham_dinh_Nghi_quyet.pdf"),
    Spec(12, "8. BC-tiep thu giai trinh BC tham dinh cua So Tu phap.pdf", "12_BC_tiep_thu_giai_trinh_tham_dinh_So_Tu_phap.pdf"),
    Spec(13, "9. Bao cao tham tra NQ thu hut nhan luc chat luong sua.signed.pdf", "13_Bao_cao_tham_tra_Nghi_quyet_HDND.pdf"),
    Spec(14, "BAO CAO TONG HOP Y KIEN TVUBND TINH DOI VOI DU THAO NQ.1.signed.signed.signed.pdf", "14_BC_tong_hop_y_kien_Thanh_vien_UBND_tinh.pdf"),
    Spec(15, "BC-giai trinh y kien cua thanh vien UBND tinh G.pdf", "15_BC_giai_trinh_y_kien_Thanh_vien_UBND_tinh.pdf"),
    Spec(16, "TTr-trinh ky thong qua lai du thao NQ (1) (1).pdf", "16_TTr_trinh_ky_thong_qua_lai_du_thao_NQ.pdf"),
    Spec(17, "Vv thong bao ket qua dang tai du thao Nghi quyet ban hanh quy dinh chinh sach ho tro dao tao- boi duong CBCCVC.signed.pdf", "17_Thong_bao_ket_qua_dang_tai_du_thao_Nghi_quyet.pdf"),
    Spec(18, "811 BC-UBND.signed.signed-1.pdf", "18_811_BC_UBND_tiep_thu_giai_trinh_tham_tra.pdf"),
]


def norm(value: str) -> str:
    value = unicodedata.normalize("NFD", value or "")
    value = "".join(ch for ch in value if unicodedata.category(ch) != "Mn")
    value = value.lower().replace("đ", "d")
    value = re.sub(r"signed", "", value)
    value = re.sub(r"[^a-z0-9]+", " ", value)
    return " ".join(value.split())


def score(target: str, candidate: str) -> float:
    a, b = norm(target), norm(candidate)
    if not a or not b:
        return 0.0
    if a in b or b in a:
        return 0.96
    aset, bset = set(a.split()), set(b.split())
    token_score = len(aset & bset) / max(1, len(aset | bset))
    seq_score = SequenceMatcher(None, a, b).ratio()
    return 0.62 * token_score + 0.38 * seq_score


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def extract_text(path: Path) -> str:
    try:
        if path.suffix.lower() == ".pdf":
            reader = PdfReader(str(path))
            parts = []
            for i, page in enumerate(reader.pages, 1):
                txt = page.extract_text() or ""
                parts.append(f"\n--- TRANG {i} ---\n{txt.strip()}\n")
            return "".join(parts).strip()
        if path.suffix.lower() == ".docx":
            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text for cell in row.cells))
            return "\n".join(parts).strip()
    except Exception as exc:  # keep sync robust even if one text extraction fails
        return f"[Không thể trích xuất text tự động: {exc}]"
    return ""


def all_link_records(page: Page) -> list[dict]:
    return page.locator("a").evaluate_all(
        "els => els.map((a,i)=>({i,text:(a.innerText||a.textContent||'').trim(),href:a.href||'',onclick:a.getAttribute('onclick')||''}))"
    )


def click_agenda_item(page: Page) -> None:
    # Attachment table can be hidden until the agenda item is selected.
    patterns = ["617/TTr-UBND", "nguồn nhân lực chất lượng cao", "nguon nhan luc chat luong cao"]
    for pattern in patterns:
        try:
            loc = page.get_by_text(re.compile(re.escape(pattern), re.I)).first
            if loc.count():
                loc.scroll_into_view_if_needed(timeout=5000)
                loc.click(timeout=7000)
                page.wait_for_timeout(1800)
                return
        except Exception:
            pass


def best_link(page: Page, target_name: str, used: set[int]) -> tuple[dict | None, float]:
    best, best_score = None, 0.0
    for rec in all_link_records(page):
        if rec["i"] in used:
            continue
        candidate = " ".join([rec.get("text", ""), rec.get("href", ""), rec.get("onclick", "")])
        s = score(target_name, candidate)
        # Strong boost for matching extension and document-number prefix.
        ext = Path(target_name).suffix.lower()
        if ext and ext in candidate.lower():
            s += 0.05
        if s > best_score:
            best, best_score = rec, s
    return best, min(best_score, 1.0)


def save_via_request(page: Page, url: str, dest: Path) -> bool:
    try:
        response = page.request.get(url, timeout=60000)
        if not response.ok:
            return False
        body = response.body()
        if len(body) < 1000:
            return False
        dest.write_bytes(body)
        return True
    except Exception:
        return False


def save_via_click(page: Page, link_index: int, dest: Path) -> bool:
    loc = page.locator("a").nth(link_index)
    try:
        with page.expect_download(timeout=30000) as info:
            loc.click(timeout=10000, force=True)
        download = info.value
        download.save_as(str(dest))
        return dest.exists() and dest.stat().st_size > 1000
    except PlaywrightTimeoutError:
        return False
    except Exception:
        return False


def main() -> int:
    FILES_DIR.mkdir(parents=True, exist_ok=True)
    TEXT_DIR.mkdir(parents=True, exist_ok=True)

    results = []
    failures = []

    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True)
        context = browser.new_context(
            accept_downloads=True,
            locale="vi-VN",
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/150 Safari/537.36",
        )
        page = context.new_page()
        page.goto(MEETING_URL, wait_until="domcontentloaded", timeout=90000)
        page.wait_for_timeout(2000)

        # First attempt: the filenames may already exist in the DOM.
        body_before = norm(page.locator("body").inner_text())
        if "617 ttr ubnd" in body_before and "811 bc ubnd" not in body_before:
            click_agenda_item(page)
        elif "811 bc ubnd" not in body_before:
            click_agenda_item(page)

        # Sometimes clicking opens an overlay/modal. Give it time to populate links.
        page.wait_for_timeout(1500)
        used_links: set[int] = set()

        for spec in SPECS:
            dest = FILES_DIR / spec.output_name
            rec, match_score = best_link(page, spec.source_name, used_links)
            success = False
            source_url = ""

            if rec and match_score >= 0.43:
                used_links.add(rec["i"])
                href = rec.get("href") or ""
                if href and not href.lower().startswith("javascript:"):
                    source_url = href
                    success = save_via_request(page, urljoin(page.url, href), dest)
                if not success:
                    success = save_via_click(page, rec["i"], dest)

            # Fallback: fuzzy-search any visible element text, then click nearest anchor.
            if not success:
                target_norm = norm(spec.source_name)
                candidates = page.locator("a,button,[role=button],td,span,div").evaluate_all(
                    "els => els.map((e,i)=>({i,text:(e.innerText||e.textContent||'').trim(),tag:e.tagName}))"
                )
                ranked = sorted(((score(spec.source_name, c["text"]), c) for c in candidates), reverse=True, key=lambda x: x[0])
                for s, cand in ranked[:8]:
                    if s < 0.58:
                        break
                    try:
                        locator = page.locator("a,button,[role=button],td,span,div").nth(cand["i"])
                        anchor = locator.locator("xpath=ancestor-or-self::a[1]")
                        clickable = anchor if anchor.count() else locator
                        with page.expect_download(timeout=15000) as info:
                            clickable.click(force=True, timeout=8000)
                        info.value.save_as(str(dest))
                        success = dest.exists() and dest.stat().st_size > 1000
                        if success:
                            break
                    except Exception:
                        continue

            if not success:
                failures.append({"index": spec.index, "name": spec.source_name, "best_match_score": round(match_score, 3)})
                continue

            text = extract_text(dest)
            text_path = TEXT_DIR / f"{spec.index:02d}.txt"
            header = (
                f"TÀI LIỆU {spec.index:02d}/18\n"
                f"Tên nguồn: {spec.source_name}\n"
                f"Tệp lưu: {spec.output_name}\n"
                f"Nguồn kỳ họp: {MEETING_URL}\n"
                f"SHA-256: {sha256(dest)}\n"
                f"{'=' * 78}\n\n"
            )
            text_path.write_text(header + text + "\n", encoding="utf-8")
            results.append(
                {
                    "index": spec.index,
                    "source_name": spec.source_name,
                    "file": f"files/{spec.output_name}",
                    "text": f"text/{spec.index:02d}.txt",
                    "bytes": dest.stat().st_size,
                    "sha256": sha256(dest),
                    "matched_url": source_url,
                    "match_score": round(match_score, 3),
                }
            )

        browser.close()

    report = {
        "meeting_url": MEETING_URL,
        "expected": len(SPECS),
        "downloaded": len(results),
        "results": results,
        "failures": failures,
    }
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    print(json.dumps(report, ensure_ascii=False, indent=2))
    if failures:
        print(f"ERROR: downloaded {len(results)}/{len(SPECS)} documents", file=sys.stderr)
        return 2
    print("OK: synchronized all 18 official documents")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
